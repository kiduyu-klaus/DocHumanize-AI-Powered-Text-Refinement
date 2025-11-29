#!/usr/bin/env python3
"""
DocHumanize - AI-Powered Text Refinement
Main entry point with threading support for faster processing
"""

import argparse
import os
import sys
from concurrent.futures import ThreadPoolExecutor, as_completed
from threading import Lock
from docx import Document
from docx_processor import apply_run_format, process_docx, batch_process_docx
from ollama_humanize import humanize_with_ollama


# Thread-safe print lock
print_lock = Lock()


def safe_print(*args, **kwargs):
    """Thread-safe print function"""
    with print_lock:
        print(*args, **kwargs)


def process_paragraph_threaded(paragraph_data):
    """
    Process a single paragraph in a thread.
    
    :param paragraph_data: tuple of (index, paragraph, ollama_model, ollama_url, temperature, max_tokens, preserve_formatting)
    :return: tuple of (index, success, humanized_text, formatting_data, error_message)
    """
    index, paragraph, ollama_model, ollama_url, temperature, max_tokens, preserve_formatting = paragraph_data
    
    original_text = paragraph.text
    
    try:
        humanized_text = humanize_with_ollama(
            original_text,
            model=ollama_model,
            ollama_url=ollama_url,
            temperature=temperature,
            max_tokens=max_tokens
        )
        
        formatting_data = None
        if preserve_formatting:
            # Store formatting information
            run_formats = []
            for run in paragraph.runs:
                run_formats.append({
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    'font_name': run.font.name,
                    'font_size': run.font.size,
                    'font_color': run.font.color.rgb if run.font.color.rgb else None
                })
            
            paragraph_format = {
                'alignment': paragraph.alignment,
                'left_indent': paragraph.paragraph_format.left_indent,
                'right_indent': paragraph.paragraph_format.right_indent,
                'first_line_indent': paragraph.paragraph_format.first_line_indent,
                'space_before': paragraph.paragraph_format.space_before,
                'space_after': paragraph.paragraph_format.space_after,
                'line_spacing': paragraph.paragraph_format.line_spacing
            }
            
            formatting_data = {
                'run_formats': run_formats,
                'paragraph_format': paragraph_format
            }
        
        return (index, True, humanized_text, formatting_data, None)
        
    except Exception as e:
        return (index, False, None, None, str(e))


def process_docx_threaded(input_path, ollama_model="cogito-2.1:671b-cloud", 
                         ollama_url="http://localhost:11434", temperature=0.7, 
                         max_tokens=2000, preserve_formatting=True, max_workers=4):
    """
    Process a docx file using multiple threads for parallel paragraph processing.
    
    :param input_path: path to the input docx file
    :param ollama_model: the Ollama model to use
    :param ollama_url: the URL of the Ollama API
    :param temperature: controls randomness
    :param max_tokens: maximum tokens per request
    :param preserve_formatting: whether to preserve formatting
    :param max_workers: number of threads to use (default: 4)
    :return: path to the output file
    """
    
    # Validate input file
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"Input file not found: {input_path}")
    
    if not input_path.endswith('.docx'):
        raise ValueError("Input file must be a .docx file")
    
    # Generate output path
    directory = os.path.dirname(input_path) or '.'
    filename = os.path.basename(input_path)
    name, ext = os.path.splitext(filename)
    output_path = os.path.join(directory, f"{name}_edited{ext}")
    
    safe_print(f"Reading document: {input_path}")
    
    # Load the document
    doc = Document(input_path)
    
    # Collect all non-empty paragraphs
    paragraphs_to_process = []
    paragraph_indices = []
    
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip():
            paragraphs_to_process.append((
                i, paragraph, ollama_model, ollama_url, 
                temperature, max_tokens, preserve_formatting
            ))
            paragraph_indices.append(i)
    
    total_paragraphs = len(paragraphs_to_process)
    safe_print(f"Processing {total_paragraphs} paragraphs using {max_workers} threads...")
    
    # Process paragraphs in parallel
    results = {}
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        # Submit all tasks
        future_to_index = {
            executor.submit(process_paragraph_threaded, para_data): para_data[0] 
            for para_data in paragraphs_to_process
        }
        
        # Process completed tasks
        completed = 0
        for future in as_completed(future_to_index):
            index, success, humanized_text, formatting_data, error = future.result()
            results[index] = (success, humanized_text, formatting_data, error)
            
            completed += 1
            if success:
                safe_print(f"✓ Processed paragraph {completed}/{total_paragraphs} (index: {index})")
            else:
                safe_print(f"✗ Failed paragraph {completed}/{total_paragraphs} (index: {index}): {error}")
    
    # Apply results back to document in order
    safe_print("Applying changes to document...")
    success_count = 0
    
    for i, paragraph in enumerate(doc.paragraphs):
        if i not in results:
            continue
        
        success, humanized_text, formatting_data, error = results[i]
        
        if not success:
            safe_print(f"Warning: Keeping original text for paragraph {i}: {error}")
            continue
        
        # Apply the humanized text
        if preserve_formatting and formatting_data:
            # Clear all runs
            for run in paragraph.runs:
                run.text = ''
            
            # Add humanized text to first run or create new one
            if paragraph.runs:
                paragraph.runs[0].text = humanized_text
                if formatting_data['run_formats']:
                    apply_run_format(paragraph.runs[0], formatting_data['run_formats'][0])
            else:
                run = paragraph.add_run(humanized_text)
                if formatting_data['run_formats']:
                    apply_run_format(run, formatting_data['run_formats'][0])
            
            # Restore paragraph formatting
            pf = formatting_data['paragraph_format']
            paragraph.alignment = pf['alignment']
            paragraph.paragraph_format.left_indent = pf['left_indent']
            paragraph.paragraph_format.right_indent = pf['right_indent']
            paragraph.paragraph_format.first_line_indent = pf['first_line_indent']
            paragraph.paragraph_format.space_before = pf['space_before']
            paragraph.paragraph_format.space_after = pf['space_after']
            paragraph.paragraph_format.line_spacing = pf['line_spacing']
        else:
            paragraph.text = humanized_text
        
        success_count += 1
    
    # Process tables (sequential for simplicity - tables are usually smaller)
    safe_print("Processing tables...")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if not paragraph.text.strip():
                        continue
                    
                    try:
                        humanized_text = humanize_with_ollama(
                            paragraph.text,
                            model=ollama_model,
                            ollama_url=ollama_url,
                            temperature=temperature,
                            max_tokens=max_tokens
                        )
                        paragraph.text = humanized_text
                    except Exception as e:
                        safe_print(f"Warning: Could not process table cell: {str(e)}")
    
    # Save the document
    safe_print(f"Saving edited document: {output_path}")
    doc.save(output_path)
    safe_print(f"Done! Successfully processed {success_count}/{total_paragraphs} paragraphs.")
    
    return output_path


def batch_process_threaded(directory, file_pattern="*.docx", ollama_model="cogito-2.1:671b-cloud",
                           ollama_url="http://localhost:11434", temperature=0.7,
                           max_tokens=2000, preserve_formatting=True, max_workers=4):
    """
    Process multiple docx files with threading.
    
    :param directory: directory containing docx files
    :param file_pattern: pattern to match files
    :param ollama_model: the Ollama model to use
    :param ollama_url: the URL of the Ollama API
    :param temperature: controls randomness
    :param max_tokens: maximum tokens per request
    :param preserve_formatting: whether to preserve formatting
    :param max_workers: number of threads per file
    :return: list of output file paths
    """
    import glob
    
    if not os.path.exists(directory):
        raise FileNotFoundError(f"Directory not found: {directory}")
    
    pattern = os.path.join(directory, file_pattern)
    files = [f for f in glob.glob(pattern) if not f.endswith('_edited.docx')]
    
    if not files:
        safe_print(f"No docx files found in {directory}")
        return []
    
    output_files = []
    
    for i, input_file in enumerate(files):
        safe_print(f"\n{'=' * 80}")
        safe_print(f"Processing file {i + 1}/{len(files)}: {os.path.basename(input_file)}")
        safe_print('=' * 80)
        
        try:
            output_file = process_docx_threaded(
                input_file,
                ollama_model=ollama_model,
                ollama_url=ollama_url,
                temperature=temperature,
                max_tokens=max_tokens,
                preserve_formatting=preserve_formatting,
                max_workers=max_workers
            )
            output_files.append(output_file)
        except Exception as e:
            safe_print(f"Error processing {input_file}: {str(e)}")
            continue
    
    safe_print(f"\n{'=' * 80}")
    safe_print(f"Batch processing complete! Processed {len(output_files)}/{len(files)} files.")
    safe_print('=' * 80)
    
    return output_files


def main():
    """Main CLI entry point"""
    parser = argparse.ArgumentParser(
        description="DocHumanize - AI-Powered Text Refinement",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Process a single document with default settings
  python main.py --input document.docx

  # Process with custom model and 8 threads
  python main.py --input document.docx --model llama2 --threads 8

  # Batch process all docx files in a directory
  python main.py --batch-dir ./documents --threads 4

  # Custom output path and style
  python main.py --input doc.docx --output humanized.docx --temperature 0.9
        """
    )
    
    # Input options
    input_group = parser.add_mutually_exclusive_group(required=True)
    input_group.add_argument('--input', '-i', help='Input docx file path')
    input_group.add_argument('--batch-dir', '-b', help='Directory for batch processing')
    
    # Output options
    parser.add_argument('--output', '-o', help='Output file path (only for single file mode)')
    
    # Model options
    parser.add_argument('--model', '-m', default='cogito-2.1:671b-cloud',
                       help='Ollama model to use (default: cogito-2.1:671b-cloud)')
    parser.add_argument('--url', default='http://localhost:11434',
                       help='Ollama API URL (default: http://localhost:11434)')
    parser.add_argument('--temperature', '-t', type=float, default=0.7,
                       help='Temperature for text generation (0.0-1.0, default: 0.7)')
    parser.add_argument('--max-tokens', type=int, default=2000,
                       help='Maximum tokens per request (default: 2000)')
    
    # Processing options
    parser.add_argument('--threads', type=int, default=4,
                       help='Number of threads to use (default: 4)')
    parser.add_argument('--no-preserve-formatting', action='store_true',
                       help='Do not preserve document formatting')
    parser.add_argument('--no-threading', action='store_true',
                       help='Use sequential processing instead of threading')
    
    args = parser.parse_args()
    
    preserve_formatting = not args.no_preserve_formatting
    
    try:
        if args.input:
            # Single file mode
            if args.no_threading:
                output_path = process_docx(
                    args.input,
                    ollama_model=args.model,
                    ollama_url=args.url,
                    temperature=args.temperature,
                    max_tokens=args.max_tokens,
                    preserve_formatting=preserve_formatting
                )
            else:
                output_path = process_docx_threaded(
                    args.input,
                    ollama_model=args.model,
                    ollama_url=args.url,
                    temperature=args.temperature,
                    max_tokens=args.max_tokens,
                    preserve_formatting=preserve_formatting,
                    max_workers=args.threads
                )
            
            # Handle custom output path
            if args.output and output_path != args.output:
                import shutil
                shutil.move(output_path, args.output)
                print(f"Moved output to: {args.output}")
        
        else:
            # Batch mode
            if args.no_threading:
                batch_process_docx(
                    args.batch_dir,
                    ollama_model=args.model,
                    ollama_url=args.url,
                    temperature=args.temperature,
                    max_tokens=args.max_tokens,
                    preserve_formatting=preserve_formatting
                )
            else:
                batch_process_threaded(
                    args.batch_dir,
                    ollama_model=args.model,
                    ollama_url=args.url,
                    temperature=args.temperature,
                    max_tokens=args.max_tokens,
                    preserve_formatting=preserve_formatting,
                    max_workers=args.threads
                )
    
    except KeyboardInterrupt:
        print("\n\nProcess interrupted by user.")
        sys.exit(1)
    except Exception as e:
        print(f"\nError: {str(e)}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()