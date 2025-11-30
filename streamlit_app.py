import streamlit as st
import os
import tempfile
import shutil
from pathlib import Path
import time
import re
from docx import Document
from docx_processor import process_docx_with_progress
from ollama_humanize import humanize_with_ollama
import requests

# Page configuration
st.set_page_config(
    page_title="DocHumanize - AI Text Refinement",
    page_icon="📝",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for better styling
st.markdown("""
<style>
    .main-header {
        font-size: 3rem;
        font-weight: 700;
        color: #1f77b4;
        text-align: center;
        margin-bottom: 0.5rem;
    }
    .sub-header {
        font-size: 1.2rem;
        color: #666;
        text-align: center;
        margin-bottom: 2rem;
    }
    .stProgress > div > div > div > div {
        background-color: #1f77b4;
    }
    .success-box {
        padding: 1rem;
        border-radius: 0.5rem;
        background-color: #d4edda;
        border: 1px solid #c3e6cb;
        color: #155724;
        margin: 1rem 0;
    }
    .info-box {
        padding: 1rem;
        border-radius: 0.5rem;
        background-color: #d1ecf1;
        border: 1px solid #bee5eb;
        color: #0c5460;
        margin: 1rem 0;
    }
    .warning-box {
        padding: 1rem;
        border-radius: 0.5rem;
        background-color: #fff3cd;
        border: 1px solid #ffeaa7;
        color: #856404;
        margin: 1rem 0;
    }
</style>
""", unsafe_allow_html=True)


def sent_tokenize(text):
    """Simple sentence tokenizer"""
    sentences = re.split(r'[.!?]+', text)
    return [s.strip() for s in sentences if s.strip()]


def calculate_humanness_score(text):
    """
    Calculate human-likeness score based on various text characteristics
    
    :param text: the text to analyze
    :return: tuple of (score, metrics_dict)
    """
    humanness_score = 0
    metrics = {}
    
    # Sentence analysis
    sentences = sent_tokenize(text)
    sentence_count = len(sentences)
    metrics['sentence_count'] = sentence_count
    
    if sentence_count == 0:
        return 0, metrics
    
    # Sentence length variance
    sent_lengths = [len(s.split()) for s in sentences]
    if sent_lengths:
        avg_sentence_length = sum(sent_lengths) / len(sent_lengths)
        sentence_length_variance = sum((x - avg_sentence_length) ** 2 for x in sent_lengths) / len(sent_lengths)
        metrics['avg_sentence_length'] = round(avg_sentence_length, 1)
        metrics['sentence_variance'] = round(sentence_length_variance, 1)
        
        if sentence_length_variance > 10:
            humanness_score += 20
        elif sentence_length_variance > 5:
            humanness_score += 10
    
    # Contractions (natural human writing)
    contractions = len(re.findall(r"\b\w+'[a-z]+\b", text))
    metrics['contractions'] = contractions
    if contractions > 0:
        humanness_score += min(15, contractions * 3)
    
    # Transition words (good variety indicates natural flow)
    transitions = len(re.findall(
        r'\b(however|nevertheless|therefore|thus|consequently|furthermore|moreover|'
        r'in addition|in fact|actually|basically|arguably|indeed|instead|meanwhile|'
        r'nonetheless|otherwise|likewise|similarly|in other words|for example|'
        r'for instance|in particular|specifically|especially|notably|chiefly|mainly|mostly)\b',
        text.lower()
    ))
    metrics['transitions'] = transitions
    humanness_score += min(15, transitions * 3)
    
    # Filler words (very human, but not too many)
    fillers = len(re.findall(
        r'\b(um|uh|like|you know|sort of|kind of|literally|basically|actually|'
        r'anyway|so|well|right|okay|just)\b',
        text.lower()
    ))
    metrics['fillers'] = fillers
    humanness_score += min(10, fillers * 2)
    
    # Penalty for monotonous sentence length (robotic)
    if sentence_count > 5 and sent_lengths:
        if abs(max(sent_lengths) - min(sent_lengths)) < 3:
            humanness_score -= 20
            metrics['monotonous'] = True
        else:
            metrics['monotonous'] = False
    
    # Check for repeated phrases (AI tends to repeat)
    words = text.lower().split()
    three_grams = [' '.join(words[i:i+3]) for i in range(len(words)-2)]
    repeated_phrases = len(three_grams) - len(set(three_grams))
    metrics['repeated_phrases'] = repeated_phrases
    if repeated_phrases > 3:
        humanness_score -= min(20, repeated_phrases * 2)
    
    # Normalize to 0-100 scale
    humanness_score = max(0, min(100, humanness_score + 50))
    
    return humanness_score, metrics


def check_ollama_connection(url):
    """Check if Ollama is running and accessible"""
    try:
        response = requests.get(f"{url}/api/tags", timeout=5)
        return response.status_code == 200
    except:
        return False


def get_available_models(url):
    """Get list of available Ollama models"""
    try:
        response = requests.get(f"{url}/api/tags", timeout=5)
        if response.status_code == 200:
            models = response.json().get("models", [])
            return [model["name"] for model in models]
        return []
    except:
        return []


def extract_text_from_docx(file_path):
    """Extract all text from a docx file"""
    doc = Document(file_path)
    text = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text.append(paragraph.text)
    return '\n\n'.join(text)


def process_file_with_progress(uploaded_file, model, url, temperature, max_tokens, preserve_formatting):
    """Process uploaded file with progress tracking"""
    
    # Create temporary directory
    with tempfile.TemporaryDirectory() as temp_dir:
        # Save uploaded file
        input_path = os.path.join(temp_dir, uploaded_file.name)
        with open(input_path, "wb") as f:
            f.write(uploaded_file.getbuffer())
        
        # Extract original text for comparison
        original_text = extract_text_from_docx(input_path)
        
        # Progress tracking
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        def progress_callback(current, total, message):
            progress = current / total if total > 0 else 0
            progress_bar.progress(progress)
            status_text.text(f"{message} ({current}/{total})")
        
        try:
            # Process the document
            output_path = process_docx_with_progress(
                input_path,
                ollama_model=model,
                ollama_url=url,
                temperature=temperature,
                max_tokens=max_tokens,
                preserve_formatting=preserve_formatting,
                progress_callback=progress_callback
            )
            
            # Extract processed text
            processed_text = extract_text_from_docx(output_path)
            
            # Read the processed file
            with open(output_path, "rb") as f:
                processed_data = f.read()
            
            progress_bar.progress(1.0)
            status_text.text("✅ Processing complete!")
            
            return processed_data, os.path.basename(output_path), original_text, processed_text
            
        except Exception as e:
            st.error(f"Error processing file: {str(e)}")
            return None, None, None, None


def main():
    # Header
    st.markdown('<h1 class="main-header">📝 DocHumanize</h1>', unsafe_allow_html=True)
    st.markdown('<p class="sub-header">Transform AI-generated text into natural, human-like writing</p>', unsafe_allow_html=True)
    
    # Sidebar configuration
    with st.sidebar:
        st.header("⚙️ Configuration")
        
        # Ollama settings
        st.subheader("Ollama Settings")
        ollama_url = st.text_input(
            "Ollama API URL",
            value="http://localhost:11434",
            help="URL where your Ollama instance is running"
        )
        
        # Check connection
        if st.button("🔍 Test Connection"):
            with st.spinner("Testing connection..."):
                if check_ollama_connection(ollama_url):
                    st.success("✅ Connected to Ollama!")
                else:
                    st.error("❌ Cannot connect to Ollama. Make sure it's running.")
        
        # Get available models
        available_models = get_available_models(ollama_url)
        
        if available_models:
            selected_model = st.selectbox(
                "Select Model",
                options=available_models,
                index=available_models.index("cogito-2.1:671b-cloud") if "cogito-2.1:671b-cloud" in available_models else 0,
                help="Choose the Ollama model to use for text humanization"
            )
        else:
            selected_model = st.text_input(
                "Model Name",
                value="cogito-2.1:671b-cloud",
                help="Enter the model name manually"
            )
            st.warning("⚠️ Could not fetch models. Ensure Ollama is running.")
        
        st.divider()
        
        # Processing settings
        st.subheader("Processing Settings")
        
        temperature = st.slider(
            "Temperature",
            min_value=0.0,
            max_value=1.0,
            value=0.7,
            step=0.1,
            help="Controls creativity: lower = more consistent, higher = more creative"
        )
        
        max_tokens = st.number_input(
            "Max Tokens",
            min_value=500,
            max_value=4000,
            value=2000,
            step=100,
            help="Maximum length of generated text per paragraph"
        )
        
        preserve_formatting = st.checkbox(
            "Preserve Formatting",
            value=True,
            help="Maintain original document formatting (bold, italic, fonts, etc.)"
        )
        
        st.divider()
        
        # Information
        st.subheader("ℹ️ About")
        st.info(
            "DocHumanize uses AI to transform your documents into more natural, "
            "human-like text while preserving the original meaning and formatting."
        )
        
        with st.expander("📖 How to Use"):
            st.markdown("""
            1. **Upload** your .docx file
            2. **Configure** settings in the sidebar
            3. **Click** "Process Document"
            4. **Download** your humanized document
            
            **Tips:**
            - Higher temperature = more creative output
            - Lower temperature = more consistent output
            - Processing time depends on document length
            """)
    
    # Main content area
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.header("📤 Upload Document")
        uploaded_file = st.file_uploader(
            "Choose a .docx file",
            type=['docx'],
            help="Upload a Microsoft Word document (.docx format)"
        )
        
        if uploaded_file:
            st.success(f"✅ File uploaded: {uploaded_file.name}")
            
            # Show file info
            file_size = len(uploaded_file.getbuffer()) / 1024
            st.info(f"📊 File size: {file_size:.2f} KB")
            
            # Preview document stats
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                tmp_file.write(uploaded_file.getbuffer())
                tmp_path = tmp_file.name
            
            try:
                doc = Document(tmp_path)
                paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])
                word_count = sum(len(p.text.split()) for p in doc.paragraphs)
                
                col_a, col_b = st.columns(2)
                with col_a:
                    st.metric("Paragraphs", paragraph_count)
                with col_b:
                    st.metric("Words (approx)", word_count)
                
                os.unlink(tmp_path)
            except Exception as e:
                st.warning(f"Could not preview document: {str(e)}")
    
    with col2:
        st.header("🎯 Quick Actions")
        
        # Process button
        process_button = st.button(
            "🚀 Process Document",
            type="primary",
            use_container_width=True,
            disabled=(uploaded_file is None)
        )
        
        if uploaded_file:
            estimated_time = len(uploaded_file.getbuffer()) / 10240  # Rough estimate
            st.caption(f"⏱️ Estimated time: ~{estimated_time:.1f} min")
    
    # Processing section
    if process_button and uploaded_file:
        st.divider()
        st.header("⚡ Processing")
        
        start_time = time.time()
        
        # Process the document
        processed_data, output_filename, original_text, processed_text = process_file_with_progress(
            uploaded_file,
            selected_model,
            ollama_url,
            temperature,
            max_tokens,
            preserve_formatting
        )
        
        if processed_data and processed_text:
            elapsed_time = time.time() - start_time
            
            st.markdown(
                f'<div class="success-box">✅ <strong>Success!</strong> '
                f'Document processed in {elapsed_time:.1f} seconds</div>',
                unsafe_allow_html=True
            )
            
            # Calculate humanness scores
            original_score, original_metrics = calculate_humanness_score(original_text)
            processed_score, processed_metrics = calculate_humanness_score(processed_text)
            
            # Human-likeness Score Analysis
            st.header("🎯 Human-likeness Analysis")
            
            col_score1, col_score2 = st.columns(2)
            
            with col_score1:
                st.subheader("Original Document")
                color_orig = 'green' if original_score > 70 else 'orange' if original_score > 40 else 'red'
                st.markdown(f"""
                <div style="text-align:center">
                    <div style="margin:20px auto; width:200px; height:200px; position:relative;">
                        <div style="position:absolute; width:200px; height:200px; border-radius:50%; background:conic-gradient(from 0deg, {color_orig} 0%, {color_orig} {original_score}%, #e0e0e0 {original_score}%, #e0e0e0 100%);"></div>
                        <div style="position:absolute; width:150px; height:150px; border-radius:50%; background:white; top:25px; left:25px; display:flex; align-items:center; justify-content:center;">
                            <span style="font-size:40px; font-weight:bold; color:black;">{original_score}%</span>
                        </div>
                    </div>
                </div>
                """, unsafe_allow_html=True)
                
                with st.expander("📊 Detailed Metrics"):
                    st.write(f"**Sentences:** {original_metrics.get('sentence_count', 0)}")
                    st.write(f"**Avg Sentence Length:** {original_metrics.get('avg_sentence_length', 0)} words")
                    st.write(f"**Sentence Variance:** {original_metrics.get('sentence_variance', 0)}")
                    st.write(f"**Contractions:** {original_metrics.get('contractions', 0)}")
                    st.write(f"**Transitions:** {original_metrics.get('transitions', 0)}")
                    st.write(f"**Fillers:** {original_metrics.get('fillers', 0)}")
                    st.write(f"**Repeated Phrases:** {original_metrics.get('repeated_phrases', 0)}")
            
            with col_score2:
                st.subheader("Processed Document")
                color_proc = 'green' if processed_score > 70 else 'orange' if processed_score > 40 else 'red'
                st.markdown(f"""
                <div style="text-align:center">
                    <div style="margin:20px auto; width:200px; height:200px; position:relative;">
                        <div style="position:absolute; width:200px; height:200px; border-radius:50%; background:conic-gradient(from 0deg, {color_proc} 0%, {color_proc} {processed_score}%, #e0e0e0 {processed_score}%, #e0e0e0 100%);"></div>
                        <div style="position:absolute; width:150px; height:150px; border-radius:50%; background:white; top:25px; left:25px; display:flex; align-items:center; justify-content:center;">
                            <span style="font-size:40px; font-weight:bold; color:black;">{processed_score}%</span>
                        </div>
                    </div>
                </div>
                """, unsafe_allow_html=True)
                
                with st.expander("📊 Detailed Metrics"):
                    st.write(f"**Sentences:** {processed_metrics.get('sentence_count', 0)}")
                    st.write(f"**Avg Sentence Length:** {processed_metrics.get('avg_sentence_length', 0)} words")
                    st.write(f"**Sentence Variance:** {processed_metrics.get('sentence_variance', 0)}")
                    st.write(f"**Contractions:** {processed_metrics.get('contractions', 0)}")
                    st.write(f"**Transitions:** {processed_metrics.get('transitions', 0)}")
                    st.write(f"**Fillers:** {processed_metrics.get('fillers', 0)}")
                    st.write(f"**Repeated Phrases:** {processed_metrics.get('repeated_phrases', 0)}")
            
            # Score improvement
            improvement = processed_score - original_score
            if improvement > 0:
                st.success(f"🎉 Human-likeness improved by {improvement} points!")
            elif improvement < 0:
                st.warning(f"⚠️ Human-likeness decreased by {abs(improvement)} points. Try adjusting temperature.")
            else:
                st.info("ℹ️ Human-likeness score remained the same.")
            
            st.divider()
            
            # Download button
            st.download_button(
                label="📥 Download Processed Document",
                data=processed_data,
                file_name=output_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary",
                use_container_width=True
            )
            
            # Show text comparison
            with st.expander("📊 View Text Comparison"):
                col_before, col_after = st.columns(2)
                
                with col_before:
                    st.subheader("Original")
                    # Show first 500 characters
                    preview_orig = original_text[:500] + ("..." if len(original_text) > 500 else "")
                    st.text_area(
                        "Before",
                        value=preview_orig,
                        height=300,
                        disabled=True,
                        label_visibility="collapsed"
                    )
                
                with col_after:
                    st.subheader("Humanized")
                    # Show first 500 characters
                    preview_proc = processed_text[:500] + ("..." if len(processed_text) > 500 else "")
                    st.text_area(
                        "After",
                        value=preview_proc,
                        height=300,
                        disabled=True,
                        label_visibility="collapsed"
                    )
    
    # Footer
    st.divider()
    col_footer1, col_footer2, col_footer3 = st.columns(3)
    
    with col_footer1:
        st.caption("🔒 Your documents are processed locally")
    with col_footer2:
        st.caption("⚡ Powered by Ollama AI")
    with col_footer3:
        st.caption("📝 Supports .docx format")


if __name__ == "__main__":
    main()